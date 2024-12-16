from docx import Document
def AddPicture(image_path, docx_path):
    doc = Document(docx_path)
    doc.add_picture(image_path)
    doc.add_paragraph('Изображение успешно сохранено')
    doc.save('text.docx')

AddPicture('cars.jpg', 'text.docx')